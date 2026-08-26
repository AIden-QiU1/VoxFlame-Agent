#!/usr/bin/env python3
"""Generate the two final RanYan hardware product DOCX documents."""

from __future__ import annotations

import re
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCUMENTS = (
    (
        ROOT / "research/product-engineering/RANYAN_HARDWARE_PRODUCT_REQUIREMENTS_FIRST_SUPPLIER_2026-08-17.md",
        ROOT / "燃言多模态AI无障碍沟通机_产品功能需求文档_首家供应商修订版_2026-08-17.docx",
        "燃言多模态AI无障碍沟通机 产品功能需求文档（首家供应商修订版）",
    ),
    (
        ROOT / "research/product-engineering/RANYAN_HARDWARE_PRODUCT_PLAN_GENERAL_SUPPLIER_2026-08-17.md",
        ROOT / "燃言多模态AI无障碍沟通机_产品方案_通用供应商版_2026-08-17.docx",
        "燃言多模态AI无障碍沟通机 产品方案（通用供应商版）",
    ),
)

PACKAGE_RELATIONSHIP_NS = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)
THUMBNAIL_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/package/2006/relationships/metadata/thumbnail"
)


def font(run, name="宋体", size=10.5, bold=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def style_font(style, name: str, size: float, bold=None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    style_font(normal, "宋体", 10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    style_font(doc.styles["Title"], "黑体", 18)
    doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Title"].paragraph_format.space_after = Pt(14)
    title_ppr = doc.styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    for name, size, before, after in (
        ("Heading 1", 16, 13, 7),
        ("Heading 2", 13.5, 10, 5),
        ("Heading 3", 12, 8, 4),
    ):
        style = doc.styles[name]
        style_font(style, "黑体", size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("燃言无障碍沟通硬件产品文档　")
    font(run, size=8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def split_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def separator(line: str) -> bool:
    cells = split_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def cell_margins(cell, value=85) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_header(table.rows[0])
    for row_index, row in enumerate(rows):
        prevent_row_split(table.rows[row_index])
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(row[column_index] if column_index < len(row) else "")
            size = 8.2 if columns >= 5 else 8.8 if columns == 4 else 9.2
            font(run, size=size, bold=(row_index == 0))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_paragraph(doc: Document, text: str, numbered=False) -> None:
    paragraph = doc.add_paragraph()
    if numbered:
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    font(run)


def render(doc: Document, source: str) -> None:
    lines = source.splitlines()
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style="Title")
            index += 1
            continue
        if stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped.lstrip("> "))
            font(run, "宋体", 10, True)
            index += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            doc.add_heading(heading.group(2), level=min(level, 3))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and separator(lines[index + 1]):
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            add_paragraph(doc, f"{numbered.group(1)}. {numbered.group(2)}", numbered=True)
            index += 1
            continue
        add_paragraph(doc, stripped)
        index += 1


def remove_blank_package_thumbnail(output_path: Path) -> None:
    """Remove python-docx's blank template thumbnail from the DOCX package."""
    with tempfile.NamedTemporaryFile(
        dir=output_path.parent,
        prefix=f".{output_path.stem}-",
        suffix=".docx",
        delete=False,
    ) as temporary_file:
        temporary_path = Path(temporary_file.name)

    try:
        with zipfile.ZipFile(output_path, "r") as source_zip:
            with zipfile.ZipFile(temporary_path, "w") as target_zip:
                for item in source_zip.infolist():
                    if item.filename == "docProps/thumbnail.jpeg":
                        continue
                    payload = source_zip.read(item.filename)
                    if item.filename == "_rels/.rels":
                        root = ElementTree.fromstring(payload)
                        for relationship in list(root):
                            if relationship.get("Type") == THUMBNAIL_RELATIONSHIP_TYPE:
                                root.remove(relationship)
                        ElementTree.register_namespace("", PACKAGE_RELATIONSHIP_NS)
                        payload = ElementTree.tostring(
                            root,
                            encoding="utf-8",
                            xml_declaration=True,
                        )
                    target_zip.writestr(item, payload)
        temporary_path.replace(output_path)
    finally:
        temporary_path.unlink(missing_ok=True)


def generate(source_path: Path, output_path: Path, title: str) -> None:
    doc = Document()
    configure(doc)
    properties = doc.core_properties
    properties.title = title
    properties.subject = "燃言多模态 AI 无障碍沟通硬件产品规划、工程实现与供应商协作"
    properties.author = "燃言 / 上海生声不息科技有限公司"
    properties.keywords = "燃言, 构音障碍, 无障碍沟通, AAC, 智能硬件, 产品方案"
    properties.comments = "正式对外产品文件。具体规格以双方冻结文件、样机验证和合同为准。"
    properties.modified = datetime(2026, 8, 17, 23, 0, 0)
    render(doc, source_path.read_text(encoding="utf-8"))
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    doc.save(output_path)
    remove_blank_package_thumbnail(output_path)
    print(output_path)


def main() -> None:
    for source, output, title in DOCUMENTS:
        if not source.exists():
            raise SystemExit(f"Missing source: {source}")
        generate(source, output, title)


if __name__ == "__main__":
    main()
