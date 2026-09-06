from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "research" / "speech-health" / "燃言_构音障碍普通话语音采集标准化SOP_v1.1_Web对齐版.md"
OUTPUT = ROOT / "燃言_构音障碍普通话语音采集标准化SOP_v1.1_Web对齐版.docx"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def clean_text(value: str) -> str:
    return value.replace("**", "").replace("`", "")


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = clean_text(value.strip())
            if row_index == 0:
                set_cell_shading(cell, "E7E5E4")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    document.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
        if not all(set(value) <= {"-", ":"} for value in values):
            rows.append(values)
        index += 1
    return rows, index


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Title", 24, "1C1917"),
        ("Heading 1", 16, "92400E"),
        ("Heading 2", 13, "44403C"),
        ("Heading 3", 11, "57534E"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    index = 0
    title_done = False
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(line[2:].strip()).bold = True
            title_done = True
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith(">"):
            text = clean_text(line.lstrip("> "))
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if title_done else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(87, 83, 78)
        elif line.startswith("- "):
            document.add_paragraph(clean_text(line[2:]), style="List Bullet")
        elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.45)
            paragraph.paragraph_format.first_line_indent = Cm(-0.45)
            paragraph.add_run(clean_text(line))
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.add_run(clean_text(line))
        index += 1

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("燃言语音采集操作说明 · 页面使用版")

    document.core_properties.title = "燃言构音障碍普通话语音采集标准化 SOP v1.1 Web 对齐版"
    document.core_properties.subject = "Web 自助采集与机构辅助采集标准"
    document.core_properties.author = "燃言项目组"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
