#!/usr/bin/env python3
"""Build the formal DOCX for 沪浦网信安通〔2026〕267号."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "research/product-engineering/上海生声不息科技有限公司网络安全整改报告_沪浦网信安通2026_267号.md"
OUTPUT = ROOT / "上海生声不息科技有限公司网络安全整改报告（沪浦网信安通〔2026〕267号）.docx"

BODY_FONT = "Noto Serif CJK SC"
HEADING_FONT = "Noto Sans CJK SC"


def set_run_font(run, font: str, size: float, *, bold: bool = False) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), font)


def set_paragraph_layout(
    paragraph,
    *,
    first_line_chars: int = 0,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    keep_with_next: bool = False,
) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(28)
    if first_line_chars:
        fmt.first_line_indent = Pt(16 * first_line_chars)
    props = paragraph._element.get_or_add_pPr()
    if keep_with_next:
        props.append(OxmlElement("w:keepNext"))
    props.append(OxmlElement("w:widowControl"))


def add_text(paragraph, text: str, font: str, size: float, *, bold: bool = False) -> None:
    # Render Markdown bold spans without leaving literal asterisks in the report.
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        clean = clean.replace("`", "")
        set_run_font(paragraph.add_run(clean), font, size, bold=bold or is_bold)


def add_page_number(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_layout(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(paragraph.add_run("— "), BODY_FONT, 14)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), BODY_FONT)
    props.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "28")
    props.append(size)
    run.append(props)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    set_run_font(paragraph.add_run(" —"), BODY_FONT, 14)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    add_page_number(section)

    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(28)

    settings = document.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    setting = OxmlElement("w:compatSetting")
    setting.set(qn("w:name"), "compatibilityMode")
    setting.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
    setting.set(qn("w:val"), "15")
    compat.append(setting)


def build() -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "上海生声不息科技有限公司网络安全整改报告"
    document.core_properties.subject = "沪浦网信安通〔2026〕267号整改报告"
    document.core_properties.author = "上海生声不息科技有限公司"
    document.core_properties.keywords = "网络安全整改, 沪浦网信安通〔2026〕267号"

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    nonempty = [line.strip() for line in lines if line.strip()]
    signature_lines = {
        "上海生声不息科技有限公司（加盖公章）",
        "2026年8月5日",
    }

    for line in nonempty:
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            set_paragraph_layout(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                keep_with_next=True,
            )
            paragraph.paragraph_format.line_spacing = Pt(36)
            paragraph.paragraph_format.space_after = Pt(14)
            add_text(paragraph, line[2:], HEADING_FONT, 22, bold=True)
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph()
            set_paragraph_layout(paragraph, keep_with_next=True)
            paragraph.paragraph_format.space_before = Pt(10)
            add_text(paragraph, line[3:], HEADING_FONT, 16, bold=True)
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph()
            set_paragraph_layout(paragraph, keep_with_next=True)
            add_text(paragraph, line[4:], BODY_FONT, 16, bold=True)
            continue

        if line in signature_lines:
            paragraph = document.add_paragraph()
            set_paragraph_layout(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            paragraph.paragraph_format.right_indent = Cm(1.0)
            add_text(paragraph, line, BODY_FONT, 16)
            continue

        if line == "上海市浦东新区网络与信息安全信息通报中心：":
            paragraph = document.add_paragraph()
            set_paragraph_layout(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            add_text(paragraph, line, BODY_FONT, 16)
            continue

        paragraph = document.add_paragraph()
        numbered_item = bool(re.match(r"^\d+\.\s", line))
        set_paragraph_layout(
            paragraph,
            first_line_chars=2,
            alignment=(
                WD_ALIGN_PARAGRAPH.LEFT
                if numbered_item
                else WD_ALIGN_PARAGRAPH.JUSTIFY
            ),
        )
        add_text(paragraph, line, BODY_FONT, 16)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
